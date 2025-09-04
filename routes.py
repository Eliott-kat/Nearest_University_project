from app import app, db
from models import Document, DocumentStatus, AnalysisResult, HighlightedSentence
from file_utils import save_uploaded_file, extract_text_from_file, get_file_size
from auth_simple import is_logged_in, get_current_user, require_auth
from language_utils import LanguageManager
from werkzeug.exceptions import RequestEntityTooLarge
from pdf_annotation import generate_annotated_pdf_for_document
# Ajouter ces imports pour la génération de documents formatés
from docx import Document as DocxDocument
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_UNDERLINE
import tempfile
import os
import logging
from flask import render_template, request, redirect, url_for, flash, session, jsonify, send_file, abort
"""
Routes for AcadCheck with authentication system
"""
@app.route('/report-html/<int:document_id>')
@require_auth
def view_report_html(document_id):
    """Affiche le document Word original converti en HTML avec surlignement des passages détectés."""
    import mammoth
    import re
    user_id = session.get('user_id') or session.get('demo_user', {}).get('id')
    document = Document.query.filter_by(id=document_id, user_id=user_id).first_or_404()
    # Chemin du fichier Word original
    file_path = getattr(document, 'file_path', None)
    if not file_path:
        import os
        upload_dir = os.path.join(os.getcwd(), 'uploads')
        file_path = os.path.join(upload_dir, document.original_filename)
    # Conversion DOCX -> HTML
    with open(file_path, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        html = result.value  # HTML string

    # Récupérer les phrases à surligner
    highlighted_sentences = HighlightedSentence.query.filter_by(document_id=document.id).order_by(HighlightedSentence.start_position).all()
    for hs in highlighted_sentences:
        plag_score = hs.plagiarism_score if hs.plagiarism_score is not None else 0
        ai_score = hs.ai_score if hasattr(hs, 'ai_score') and hs.ai_score is not None else 0
        if plag_score > 30:
            color = "#ffb3b3"  # rouge clair
        elif ai_score > 40:
            color = "#b3c6ff"  # bleu clair
        else:
            color = "#ffff99"  # jaune
        # Remplacer la phrase par une version surlignée (tolérance espaces)
        pattern = re.escape(hs.sentence_text.strip())
        html = re.sub(pattern, f'<span class="highlighted" style="background:{color};">{hs.sentence_text.strip()}</span>', html, flags=re.IGNORECASE)

    # Ajouter un style CSS simple pour le surlignement
    style = """
    <style>
    .highlighted { padding:2px 2px; border-radius:3px; }
    </style>
    """
    return style + html

from app import app, db
"""
Routes for AcadCheck with authentication system
"""
import os
import logging
from flask import render_template, request, redirect, url_for, flash, session, jsonify, send_file, abort
from models import Document, DocumentStatus, AnalysisResult, HighlightedSentence
from file_utils import save_uploaded_file, extract_text_from_file, get_file_size
from auth_simple import is_logged_in, get_current_user, require_auth
from language_utils import LanguageManager
from werkzeug.exceptions import RequestEntityTooLarge
from pdf_annotation import generate_annotated_pdf_for_document
# Ajouter ces imports pour la génération de documents formatés
from docx import Document as DocxDocument
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_UNDERLINE
import tempfile

@app.before_request
def make_session_permanent():
    session.permanent = True

@app.context_processor
def inject_user():
    """Inject current_user for all templates"""
    return dict(user=get_current_user(), is_logged_in=is_logged_in())

def create_formatted_document_with_highlights(document, analysis_result):
    """
    Crée un document DOCX formaté avec des soulignements colorés basés sur les scores
    de plagiat et d'IA pour ressembler à l'original mais avec les annotations
    """
    try:
        from docx import Document as DocxDocument
        import difflib
        import os
        # Chemin du fichier original
        file_path = getattr(document, 'file_path', None)
        if not file_path:
            upload_dir = os.path.join(os.getcwd(), 'uploads')
            file_path = os.path.join(upload_dir, document.original_filename)
        doc = DocxDocument(file_path)

        # Récupérer les phrases à surligner
        highlighted_sentences = HighlightedSentence.query.filter_by(
            document_id=document.id
        ).order_by('sentence_index').all()
        highlight_map = {}
        for hs in highlighted_sentences:
            plag_score = hs.plagiarism_score if hs.plagiarism_score is not None else 0
            ai_score = hs.ai_score if hasattr(hs, 'ai_score') and hs.ai_score is not None else 0
            if plag_score > 30:
                highlight_map[hs.sentence_text.strip()] = RGBColor(200, 0, 0)
            elif ai_score > 40:
                highlight_map[hs.sentence_text.strip()] = RGBColor(0, 0, 200)

        # Appliquer le soulignement/couleur sur les runs concernés
        for para in doc.paragraphs:
            for run in para.runs:
                run_text = run.text.strip()
                for sent, color in highlight_map.items():
                    if run_text and difflib.SequenceMatcher(None, run_text, sent).ratio() > 0.95:
                        run.font.underline = True
                        run.font.color.rgb = color
                        break

        # Sauvegarder le document annoté (copie exacte de l'original, seul le soulignement est ajouté)
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        return temp_file.name
    except Exception as e:
        logging.error(f"Erreur création document formaté: {e}")
        return None

@app.route('/')
def index():
    """Main landing page - shows landing for non-authenticated users"""
    if is_logged_in():
        return redirect(url_for('dashboard'))
    return landing()

@app.route('/landing')
def landing():
    """Landing page for non-authenticated users"""
    return render_template('landing.html')

@app.route('/dashboard')
@require_auth
def dashboard():
    """User  dashboard with document statistics"""
    try:
        user_id = session.get('user_id') or session.get('demo_user', {}).get('id')
        
        recent_documents = Document.query.filter_by(user_id=user_id)\
            .order_by(Document.created_at.desc())\
            .limit(5).all()
        
        total_documents = Document.query.filter_by(user_id=user_id).count()
        completed_analyses = Document.query.filter_by(
            user_id=user_id, 
            status=DocumentStatus.COMPLETED
        ).count()
        processing_documents = Document.query.filter_by(
            user_id=user_id,
            status=DocumentStatus.PROCESSING
        ).count()
        
        completed_docs = Document.query.filter_by(
            user_id=user_id,
            status=DocumentStatus.COMPLETED
        ).all()
        
        if completed_docs:
            total_plagiarism = 0
            total_ai = 0
            count_plagiarism = 0
            count_ai = 0
            
            for doc in completed_docs:
                if doc.analysis_result:
                    if doc.analysis_result.plagiarism_score is not None:
                        total_plagiarism += doc.analysis_result.plagiarism_score
                        count_plagiarism += 1
                    if doc.analysis_result.ai_score is not None:
                        total_ai += doc.analysis_result.ai_score
                        count_ai += 1
            
            avg_plagiarism_score = total_plagiarism / count_plagiarism if count_plagiarism > 0 else 0
            avg_ai_score = total_ai / count_ai if count_ai > 0 else 0
        else:
            avg_plagiarism_score = 0
            avg_ai_score = 0
    
    except Exception as e:
        logging.error(f"Error loading dashboard: {e}")
        recent_documents = []
        avg_plagiarism_score = 0
        avg_ai_score = 0
        total_documents = 0
        completed_analyses = 0
        processing_documents = 0
    
    stats = {
        'total_documents': total_documents,
        'completed_analyses': completed_analyses,
        'processing_documents': processing_documents,
        'avg_plagiarism_score': round(avg_plagiarism_score, 1),
        'avg_ai_score': round(avg_ai_score, 1)
    }
    
    return render_template('dashboard.html', 
                         recent_documents=recent_documents, 
                         stats=stats)

@app.route('/upload', methods=['GET', 'POST'])
@require_auth
def upload_document():
    """Upload and submit document for analysis"""
    if request.method == 'POST':
        try:
            if 'file' not in request.files:
                flash('No file selected', 'danger')
                return redirect(request.url)

            file = request.files['file']
            if file.filename == '':
                flash('No file selected', 'danger')
                return redirect(request.url)

            file_info = save_uploaded_file(file)
            if not file_info:
                flash('Invalid file type. Please upload PDF, DOCX, or TXT files only.', 'danger')
                return redirect(request.url)

            file_path, filename = file_info
            content_type = file.content_type or 'text/plain'
            extracted_text = extract_text_from_file(file_path, content_type)
            if not extracted_text:
                flash('Could not extract text from the uploaded file.', 'danger')
                return redirect(request.url)

            document = Document()
            document.filename = filename
            document.original_filename = file.filename
            document.file_path = file_path
            document.file_size = get_file_size(file_path)
            document.content_type = content_type
            document.extracted_text = extracted_text
            user_id = session.get('user_id') or session.get('demo_user', {}).get('id', 'demo-user')
            document.user_id = user_id
            document.status = DocumentStatus.UPLOADED

            try:
                db.session.add(document)
                db.session.commit()
            except Exception as db_error:
                logging.error(f"Erreur sauvegarde document: {db_error}")
                db.session.rollback()
                flash('Erreur de base de données. Veuillez réessayer.', 'danger')
                return redirect(request.url)

            try:
                from ai_perplexity_detectgpt import fusion_plagiarism_score, ai_detection_score_optimized
                
                plag_score, plag_exact, plag_sem = fusion_plagiarism_score(extracted_text)
                result = {
                    'plagiarism': {
                        'percent': plag_score,
                        'exact': plag_exact,
                        'semantic': plag_sem,
                        'sources_found': 0
                    }
                }
                
                ai_result = ai_detection_score_optimized(extracted_text)
                result['ai_content'] = {
                    'percent': ai_result.get('score', 0),
                    'perplexity': ai_result.get('perplexity', 0),
                    'burstiness': ai_result.get('burstiness', 0),
                    'norm_ppl': ai_result.get('details', {}).get('norm_ppl', 0),
                    'norm_burstiness': ai_result.get('details', {}).get('norm_burstiness', 0)
                }
                
                word_count = len(extracted_text.split()) if extracted_text else 0
                plagiarism_percent = result['plagiarism']['percent']
                ai_percent = result['ai_content']['percent']

                identical_words = int((plagiarism_percent / 100.0) * word_count) if plagiarism_percent > 0 else 0
                ai_words = int((ai_percent / 100.0) * word_count) if ai_percent > 0 else 0

                analysis_result = AnalysisResult()
                analysis_result.document_id = document.id
                analysis_result.plagiarism_score = plagiarism_percent
                analysis_result.ai_score = ai_percent
                analysis_result.total_words = word_count
                analysis_result.identical_words = identical_words
                analysis_result.ai_words = ai_words
                analysis_result.sources_count = result['plagiarism']['sources_found']
                analysis_result.analysis_provider = 'local'
                analysis_result.raw_response = str(result)

                db.session.add(analysis_result)
                document.status = DocumentStatus.COMPLETED
                db.session.commit()

                # Create highlighted sentences for the PDF report
                try:
                    from simple_highlight_generator import create_highlights_for_document
                    logging.info(f"Creating highlighted sentences for document ID: {document.id}")
                    highlighted_sentences = create_highlights_for_document(document, analysis_result)
                    logging.info(f"Successfully created {len(highlighted_sentences)} highlighted sentences for document ID: {document.id}")
                except Exception as e:
                    logging.error(f"Error creating highlighted sentences for document ID: {document.id}: {e}")

                flash(f'✅ Document analysé avec succès! Plagiat: {plagiarism_percent}% + IA: {ai_percent}%', 'success')
                return redirect(url_for('document_history'))
                
            except Exception as analysis_error:
                logging.error(f"Erreur lors de l'analyse: {analysis_error}")
                db.session.rollback()
                flash('Une erreur est survenue lors de l\'analyse. Veuillez réessayer.', 'danger')
                return redirect(url_for('document_history'))

        except RequestEntityTooLarge:
            flash('File too large. Maximum file size is 16MB.', 'danger')
            return redirect(request.url)
        except Exception as e:
            logging.error(f"Error uploading document: {e}")
            flash('An error occurred while uploading the document.', 'danger')
            return redirect(request.url)

    return render_template('upload.html')

@app.route('/history')
@require_auth
def document_history():
    """View document submission history"""
    page = request.args.get('page', 1, type=int)
    per_page = 10  # Number of documents per page
    
    try:
        user_id = session.get('user_id') or session.get('demo_user', {}).get('id')
        
        documents = Document.query.filter_by(user_id=user_id)\
            .order_by(Document.created_at.desc())\
            .paginate(page=page, per_page=per_page, error_out=False)
            
        return render_template('history.html', documents=documents)
    except Exception as e:
        logging.error(f"Error loading document history: {e}")
        flash('Error loading document history.', 'danger')
        return render_template('history.html', documents=None)

@app.route('/report/<int:document_id>')
@require_auth
def view_report(document_id):
    """View detailed analysis report"""
    try:
        user_id = session.get('user_id') or session.get('demo_user', {}).get('id')
        
        document = Document.query.filter_by(
            id=document_id, 
            user_id=user_id
        ).first_or_404()
        
        if document.status != DocumentStatus.COMPLETED:
            flash('Analysis not yet completed for this document.', 'warning')
            return redirect(url_for('document_history'))
        
        analysis_result = AnalysisResult.query.filter_by(document_id=document.id).first()
        if not analysis_result:
            flash('No analysis results found for this document.', 'warning')
            return redirect(url_for('document_history'))
        
        # Utilise le formateur académique pour l'affichage pro
        from professional_document_formatter import format_academic_document
        highlighted_text = format_academic_document(
            document.extracted_text or "",
            analysis_result.plagiarism_score,
            analysis_result.ai_score,
            title=document.original_filename or "Document",
            author=getattr(document, "author", "Auteur inconnu"),
            institution=getattr(document, "institution", "Établissement non spécifié")
        )
        
        return render_template('report.html',
                             document=document,
                             analysis_result=analysis_result,
                             highlighted_text=highlighted_text)
                             
    except Exception as e:
        logging.error(f"Error loading report for document {document_id}: {e}")
        flash('Error loading report.', 'danger')
        return redirect(url_for('document_history'))

@app.route('/download-report/<int:document_id>')
@require_auth
def download_report(document_id):
    """Download formatted document with highlights"""
    try:
        user_id = session.get('user_id') or session.get('demo_user', {}).get('id')
        document = Document.query.filter_by(
            id=document_id, 
            user_id=user_id
        ).first_or_404()
        
        if document.status != DocumentStatus.COMPLETED:
            flash('Analysis not yet completed for this document.', 'warning')
            return redirect(url_for('document_history'))
        
        analysis_result = AnalysisResult.query.filter_by(document_id=document.id).first()
        if not analysis_result:
            flash('No analysis results found for this document.', 'warning')
            return redirect(url_for('view_report', document_id=document_id))
        
        # Check if document is PDF and generate annotated PDF
        if document.content_type == 'application/pdf':
            annotated_pdf_path = generate_annotated_pdf_for_document(document.id, document.file_path)
            if annotated_pdf_path:
                return send_file(
                    annotated_pdf_path,
                    as_attachment=True,
                    download_name=f"annotated_{document.original_filename}",
                    mimetype='application/pdf'
                )
        
        # Fallback to DOCX format for non-PDF documents
        docx_path = create_formatted_document_with_highlights(document, analysis_result)
        
        if not docx_path:
            flash('Error generating formatted document.', 'danger')
            return redirect(url_for('view_report', document_id=document_id))
        
        return send_file(
            docx_path,
            as_attachment=True,
            download_name=f"annotated_{document.original_filename}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logging.error(f"Error downloading report for document {document_id}: {e}")
        flash('Error downloading report.', 'danger')
        return redirect(url_for('document_history'))

@app.route('/view-annotated-pdf/<int:document_id>')
@require_auth
def view_annotated_pdf(document_id):
    """Serve annotated PDF for viewing in browser"""
    try:
        user_id = session.get('user_id') or session.get('demo_user', {}).get('id')
        document = Document.query.filter_by(
            id=document_id, 
            user_id=user_id
        ).first_or_404()
        
        if document.status != DocumentStatus.COMPLETED:
            flash('Analysis not yet completed for this document.', 'warning')
            return redirect(url_for('document_history'))
        
        # Generate or get annotated PDF path
        annotated_pdf_path = generate_annotated_pdf_for_document(document.id, document.file_path)
        
        if not annotated_pdf_path:
            flash('Could not generate annotated PDF.', 'danger')
            return redirect(url_for('view_report', document_id=document_id))
        
        return send_file(
            annotated_pdf_path,
            as_attachment=False,
            download_name=f"annotated_{document.original_filename}",
            mimetype='application/pdf'
        )
        
    except Exception as e:
        logging.error(f"Error serving annotated PDF for document {document_id}: {e}")
        flash('Error loading annotated PDF.', 'danger')
        return redirect(url_for('view_report', document_id=document_id))

@app.route('/logout')
def logout():
    """Logout route"""
    session.clear()
    flash('You have been logged out.', 'info')
    return redirect(url_for('index'))

@app.route('/login')  
def login():
    """Login route for local installation"""
    session['demo_user'] = {
        'id': 'demo-user',
        'username': 'demo_user',
        'email': 'demo@example.com'
    }
    flash('Logged in as demo user.', 'success')
    return redirect(url_for('dashboard'))

# Error handlers
@app.errorhandler(404)
def not_found_error(error):
    return render_template('404.html'), 404

@app.errorhandler(500)
def internal_error(error):
    db.session.rollback()
    return render_template('500.html'), 500

@app.errorhandler(413)
def request_entity_too_large(error):
    flash('File too large. Maximum file size is 16MB.', 'danger')
    return redirect(url_for('upload_document'))